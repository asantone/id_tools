# purpose: this tool merges multiple .docx files into a single .docx file
# author: Adam Santone, PhD
# date: 2023-06-08

#libraries
import glob, os                           # file utilities
import docx                               # (python-docx) for creating MS Word .docx files
from docx.enum.section import WD_ORIENT   # for Word document layout
from docx.enum.section import WD_SECTION  # for Word document layout
from docx.shared import Cm                # metric system
from natsort import natsorted             # for natural sorting order (1,2,3...n)
from docx2python import docx2python       # utility for docx content to Python object

#generate natural-sorted list of files to be processed
docList = []
docList = natsorted(glob.glob('./**/*.docx', recursive=True))
docList = sorted(docList, key=str.lower)

#print list to console
print(docList)

#set up the output document
document = docx.Document()

#stitching function
def stitch(filename):
    
    #get the input filenames, split them, and make them headers for each section
    fn = os.path.splitext(os.path.basename(filename))[0]
    document.add_heading(fn, level=3)    
 
    #create the content
    doc = docx2python(filename) 
    document.add_paragraph(doc.text)

#repeat for each input document    
for i in docList:
    stitch(i)

#create an output file called merged.docx and save it
outfile = "merged" + ".docx"
document.save(outfile)
